# -*- coding: utf-8 -*-
"""
PyQt6主界面
"""

import sys
from pathlib import Path

from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QListWidget, QListWidgetItem,
    QTextEdit, QProgressBar, QFileDialog, QMessageBox,
    QComboBox, QSplitter, QMenu, QInputDialog
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt6.QtGui import QAction, QColor, QFont

from utils.database import Database, TranscriptSegment
from utils.config import APP_NAME, VERSION, SUPPORTED_AUDIO_FORMATS
from core.pipeline import TranscriptionPipeline


class WorkerThread(QThread):
    """后台处理线程"""
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(str)  # recording_id
    error = pyqtSignal(str)
    
    def __init__(self, audio_path, language, model_name):
        super().__init__()
        self.audio_path = audio_path
        self.language = language
        self.model_name = model_name
        self.pipeline = None
    
    def run(self):
        try:
            self.pipeline = TranscriptionPipeline(self.model_name)
            recording_id = self.pipeline.process(
                self.audio_path,
                self.language,
                progress_callback=lambda p, msg: self.progress.emit(p, msg)
            )
            self.finished.emit(recording_id)
        except Exception as e:
            self.error.emit(str(e))
    
    def stop(self):
        if self.pipeline:
            self.pipeline.close()


class MainWindow(QMainWindow):
    """主窗口"""
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} v{VERSION}")
        self.setMinimumSize(1200, 800)
        
        self.db = Database()
        self.current_recording_id = None
        self.worker = None
        
        self.init_ui()
        self.load_recordings()
    
    def init_ui(self):
        """初始化界面"""
        # 中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QHBoxLayout(central_widget)
        
        # 分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)
        
        # 左侧：文件列表
        left_panel = self.create_left_panel()
        splitter.addWidget(left_panel)
        
        # 右侧：内容显示
        right_panel = self.create_right_panel()
        splitter.addWidget(right_panel)
        
        # 设置分割比例
        splitter.setSizes([300, 900])
        
        # 菜单栏
        self.create_menu()
    
    def create_left_panel(self):
        """创建左侧面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 标题
        title = QLabel("📁 录音文件")
        title.setFont(QFont("Microsoft YaHei", 12, QFont.Weight.Bold))
        layout.addWidget(title)
        
        # 上传按钮
        self.btn_upload = QPushButton("➕ 上传音频文件")
        self.btn_upload.setStyleSheet("""
            QPushButton {
                background-color: #4ECDC4;
                color: white;
                padding: 10px;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #45B7D1;
            }
        """)
        self.btn_upload.clicked.connect(self.upload_file)
        layout.addWidget(self.btn_upload)
        
        # 文件列表
        self.list_files = QListWidget()
        self.list_files.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.list_files.customContextMenuRequested.connect(self.show_file_menu)
        self.list_files.itemClicked.connect(self.on_file_selected)
        layout.addWidget(self.list_files)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        self.lbl_status = QLabel("")
        self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.lbl_status)
        
        return panel
    
    def create_right_panel(self):
        """创建右侧面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 顶部工具栏
        toolbar = QHBoxLayout()
        
        # 语言选择
        toolbar.addWidget(QLabel("语言:"))
        self.combo_language = QComboBox()
        self.combo_language.addItems(["中文", "英文", "日文", "自动检测"])
        toolbar.addWidget(self.combo_language)
        
        # 模型选择
        toolbar.addWidget(QLabel("模型:"))
        self.combo_model = QComboBox()
        self.combo_model.addItems(["base (推荐)", "tiny (快速)", "small (高精度)"])
        toolbar.addWidget(self.combo_model)
        
        toolbar.addStretch()
        
        # 导出按钮
        self.btn_export_txt = QPushButton("📄 导出TXT")
        self.btn_export_txt.clicked.connect(lambda: self.export_file("txt"))
        toolbar.addWidget(self.btn_export_txt)
        
        self.btn_export_word = QPushButton("📝 导出Word")
        self.btn_export_word.clicked.connect(lambda: self.export_file("word"))
        toolbar.addWidget(self.btn_export_word)
        
        self.btn_export_fav = QPushButton("⭐ 导出收藏")
        self.btn_export_fav.clicked.connect(self.export_favorites)
        toolbar.addWidget(self.btn_export_fav)
        
        layout.addLayout(toolbar)
        
        # 内容区域
        self.text_content = QTextEdit()
        self.text_content.setReadOnly(True)
        self.text_content.setStyleSheet("""
            QTextEdit {
                font-size: 14px;
                line-height: 1.6;
                padding: 10px;
            }
        """)
        layout.addWidget(self.text_content)
        
        # 信息显示
        self.lbl_info = QLabel("请上传音频文件开始转写")
        self.lbl_info.setStyleSheet("color: gray; padding: 5px;")
        layout.addWidget(self.lbl_info)
        
        return panel
    
    def create_menu(self):
        """创建菜单"""
        menubar = self.menuBar()
        
        # 文件菜单
        file_menu = menubar.addMenu("文件")
        
        open_action = QAction("打开音频", self)
        open_action.triggered.connect(self.upload_file)
        file_menu.addAction(open_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("退出", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
    
    def upload_file(self):
        """上传文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择音频文件",
            "",
            "音频文件 (*.mp3 *.wav *.m4a *.flac *.ogg *.aac *.wma)"
        )
        
        if not file_path:
            return
        
        # 获取语言和模型设置
        lang_map = {"中文": "zh", "英文": "en", "日文": "ja", "自动检测": "auto"}
        language = lang_map[self.combo_language.currentText()]
        
        model_map = {"base (推荐)": "base", "tiny (快速)": "tiny", "small (高精度)": "small"}
        model_name = model_map[self.combo_model.currentText()]
        
        # 显示进度
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.lbl_status.setText("正在处理...")
        self.btn_upload.setEnabled(False)
        
        # 启动后台线程
        self.worker = WorkerThread(file_path, language, model_name)
        self.worker.progress.connect(self.on_progress)
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()
    
    def on_progress(self, value, message):
        """进度更新"""
        self.progress_bar.setValue(value)
        self.lbl_status.setText(message)
    
    def on_finished(self, recording_id):
        """处理完成"""
        self.progress_bar.setVisible(False)
        self.lbl_status.setText("完成！")
        self.btn_upload.setEnabled(True)
        
        self.load_recordings()
        
        # 选中新文件
        for i in range(self.list_files.count()):
            item = self.list_files.item(i)
            if item.data(Qt.ItemDataRole.UserRole) == recording_id:
                self.list_files.setCurrentItem(item)
                self.on_file_selected(item)
                break
        
        QMessageBox.information(self, "完成", "音频处理完成！")
    
    def on_error(self, error_msg):
        """处理错误"""
        self.progress_bar.setVisible(False)
        self.lbl_status.setText("处理失败")
        self.btn_upload.setEnabled(True)
        
        QMessageBox.critical(self, "错误", f"处理失败:\n{error_msg}")
    
    def load_recordings(self):
        """加载录音列表"""
        self.list_files.clear()
        
        recordings = self.db.get_recordings()
        
        for rec in recordings:
            item = QListWidgetItem()
            item.setText(f"{rec.filename}\n{rec.status}")
            item.setData(Qt.ItemDataRole.UserRole, rec.id)
            
            # 根据状态设置颜色
            if rec.status == "completed":
                item.setForeground(QColor("#4ECDC4"))
            elif rec.status == "error":
                item.setForeground(QColor("#FF6B6B"))
            elif rec.status == "processing":
                item.setForeground(QColor("#F7DC6F"))
            
            self.list_files.addItem(item)
    
    def on_file_selected(self, item):
        """选择文件"""
        recording_id = item.data(Qt.ItemDataRole.UserRole)
        self.current_recording_id = recording_id
        
        # 加载转写内容
        transcripts = self.db.get_transcripts(recording_id)
        
        if not transcripts:
            self.text_content.setText("暂无转写内容")
            return
        
        # 构建显示文本
        html_parts = ["<h2>转写结果</h2>"]
        
        for seg in transcripts:
            time_str = f"{int(seg.start_time // 60):02d}:{int(seg.start_time % 60):02d}"
            speaker = seg.speaker_name or seg.speaker_id
            text = seg.text
            
            # 高亮标记
            bg_color = "#FFF9E6" if seg.is_favorite else "transparent"
            
            html_parts.append(f"""
                <div style="margin: 10px 0; padding: 8px; background: {bg_color}; border-left: 3px solid #4ECDC4;">
                    <span style="color: #888; font-size: 12px;">[{time_str}]</span>
                    <span style="color: #4ECDC4; font-weight: bold;">[{speaker}]</span>
                    <span>{text}</span>
                    {' ⭐' if seg.is_favorite else ''}
                </div>
            """)
        
        self.text_content.setHtml("".join(html_parts))
        
        # 更新信息
        self.lbl_info.setText(f"共 {len(transcripts)} 段 | 双击段落可标记收藏")
    
    def show_file_menu(self, position):
        """显示右键菜单"""
        item = self.list_files.itemAt(position)
        if not item:
            return
        
        menu = QMenu()
        
        delete_action = menu.addAction("🗑️ 删除")
        action = menu.exec(self.list_files.mapToGlobal(position))
        
        if action == delete_action:
            recording_id = item.data(Qt.ItemDataRole.UserRole)
            self.db.delete_recording(recording_id)
            self.load_recordings()
    
    def export_file(self, format_type):
        """导出文件"""
        if not self.current_recording_id:
            QMessageBox.warning(self, "提示", "请先选择一个文件")
            return
        
        transcripts = self.db.get_transcripts(self.current_recording_id)
        
        if format_type == "txt":
            file_path, _ = QFileDialog.getSaveFileName(
                self, "导出TXT", "转写结果.txt", "文本文件 (*.txt)"
            )
            if file_path:
                with open(file_path, "w", encoding="utf-8") as f:
                    for seg in transcripts:
                        time_str = f"{int(seg.start_time // 60):02d}:{int(seg.start_time % 60):02d}"
                        speaker = seg.speaker_name or seg.speaker_id
                        f.write(f"[{time_str}] [{speaker}] {seg.text}\n")
                QMessageBox.information(self, "完成", f"已导出到:\n{file_path}")
        
        elif format_type == "word":
            try:
                from docx import Document
                
                file_path, _ = QFileDialog.getSaveFileName(
                    self, "导出Word", "转写结果.docx", "Word文档 (*.docx)"
                )
                if file_path:
                    doc = Document()
                    doc.add_heading("转写结果", 0)
                    
                    for seg in transcripts:
                        time_str = f"{int(seg.start_time // 60):02d}:{int(seg.start_time % 60):02d}"
                        speaker = seg.speaker_name or seg.speaker_id
                        p = doc.add_paragraph()
                        p.add_run(f"[{time_str}] ").bold = True
                        p.add_run(f"[{speaker}] ").bold = True
                        p.add_run(seg.text)
                    
                    doc.save(file_path)
                    QMessageBox.information(self, "完成", f"已导出到:\n{file_path}")
            except ImportError:
                QMessageBox.warning(self, "错误", "请安装python-docx: pip install python-docx")
    
    def export_favorites(self):
        """导出收藏内容"""
        if not self.current_recording_id:
            QMessageBox.warning(self, "提示", "请先选择一个文件")
            return
        
        favorites = self.db.get_favorites(self.current_recording_id)
        
        if not favorites:
            QMessageBox.information(self, "提示", "没有收藏的内容")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出收藏", "收藏内容.txt", "文本文件 (*.txt)"
        )
        if file_path:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("=== 收藏的内容 ===\n\n")
                for seg in favorites:
                    time_str = f"{int(seg.start_time // 60):02d}:{int(seg.start_time % 60):02d}"
                    speaker = seg.speaker_name or seg.speaker_id
                    f.write(f"[{time_str}] [{speaker}]\n{seg.text}\n\n")
            QMessageBox.information(self, "完成", f"已导出到:\n{file_path}")
    
    def closeEvent(self, event):
        """关闭事件"""
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.worker.wait()
        
        self.db.close()
        event.accept()
